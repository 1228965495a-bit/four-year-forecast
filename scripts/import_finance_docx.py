from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import docx


ROOT = Path(__file__).resolve().parents[1]
SCRIPT_DIR = ROOT / "src" / "data" / "script"
BY_MAJOR = SCRIPT_DIR / "byMajor"

MAJOR_ID = "finance"
DOCX_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/Users/zhenzhen/Downloads/S05-金融学-完整四年脚本补强包 (1).docx")

SEMESTER_MAP = {
    "大一上": "y1s1",
    "大一下": "y1s2",
    "大二上": "y2s1",
    "大二下": "y2s2",
    "大三上": "y3s1",
    "大三下": "y3s2",
    "大四上": "y4s1",
    "大四下": "y4s2",
    "随机": None,
}

TYPE_MAP = {
    "主线事件": "main",
    "随机事件": "major_random",
}

GLOBAL_STAT_MAP = {
    "上头值": "obsession",
    "精神电量": "energy",
    "滤镜": "filter",
    "绩点意志": "gpaWill",
    "就业幻觉": "careerFantasy",
    "跑路冲动": "escapeImpulse",
}

MAJOR_STAT_MAP = {
    "市场敏感度": "quantPressure",
    "风险感知": "internshipCompetition",
    "证书焦虑": "certificateAnxiety",
}

ACHIEVEMENT_IDS = {
    "培养方案冷静派": "ach_finance_cultivation_plan",
    "高数补基础的人": "ach_finance_math_foundation",
    "变量意识觉醒者": "ach_finance_variable_awareness",
    "计量入门幸存者": "ach_finance_econometrics_survivor",
    "研报拆解者": "ach_finance_research_report",
    "实习海投人": "ach_finance_internship_mass_apply",
    "路线收敛者": "ach_finance_route_converged",
    "秋招复盘型选手": "ach_finance_autumn_review",
    "市场雷达型玩家": "ach_finance_market_radar",
    "风险收益重估者": "ach_finance_risk_return_repriced",
}

ROUTE_MARKER_MAP = {
    "postgraduate_track": "finance_route_postgrad",
    "stable_career_track": "finance_route_banking",
    "finance_career_track": "finance_route_fund_research",
}


def split_slash(value: str) -> list[str]:
    return [part.strip() for part in re.split(r"[/／]", value or "") if part.strip()]


def parse_delta(value: str) -> tuple[dict[str, int], dict[str, int]]:
    stats: dict[str, int] = {}
    major_stats: dict[str, int] = {}
    for piece in re.split(r"[，,、;；]", value or ""):
        piece = piece.strip()
        if not piece:
            continue
        match = re.match(r"(.+?)\s*([+-])\s*(\d+)", piece)
        if not match:
            continue
        name, sign, raw = match.groups()
        delta = int(raw) * (1 if sign == "+" else -1)
        name = name.strip()
        if name in GLOBAL_STAT_MAP:
            stats[GLOBAL_STAT_MAP[name]] = stats.get(GLOBAL_STAT_MAP[name], 0) + delta
        elif name in MAJOR_STAT_MAP:
            major_stats[MAJOR_STAT_MAP[name]] = major_stats.get(MAJOR_STAT_MAP[name], 0) + delta
    return stats, major_stats


def table_to_option(table, event_id: str, idx: int) -> dict:
    data = {}
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        key = cells[0].text.strip()
        val = cells[1].text.strip()
        if key:
            data[key] = val

    choice_id = ["a", "b", "c"][idx]
    text = data.get("按钮文案", "").strip()
    feedback = data.get("反馈文案", "").strip()
    stats, major_stats = parse_delta(data.get("数值变化", ""))
    tags = split_slash(data.get("标签", ""))
    route_markers = split_slash(data.get("路线标记", ""))
    achievement_names = tags + split_slash(data.get("成就", ""))

    flags = [f"{MAJOR_ID}_tag_{tag}" for tag in tags]
    route_add = [ROUTE_MARKER_MAP[m] for m in route_markers if m in ROUTE_MARKER_MAP]
    for marker in route_markers:
        flags.append(f"{MAJOR_ID}_route_marker_{marker}")

    achievement_ids = []
    for name in achievement_names:
        achievement_id = ACHIEVEMENT_IDS.get(name)
        if achievement_id:
            achievement_ids.append(achievement_id)
            flags.append(achievement_id)

    effects = {}
    if stats:
        effects["stats"] = stats
    if major_stats:
        effects["majorStats"] = major_stats
    if flags:
        effects["flagsAdd"] = sorted(set(flags), key=flags.index)
    if route_add:
        effects["routeAdd"] = sorted(set(route_add), key=route_add.index)
    if achievement_ids:
        effects["achievementIds"] = sorted(set(achievement_ids), key=achievement_ids.index)

    return {
        "id": choice_id,
        "choiceId": choice_id,
        "text": text,
        "feedback": feedback,
        "resultText": feedback,
        "effects": effects,
        "statChanges": stats,
        "routeChanges": route_add,
        "condition": None,
        "tagsUnlocked": tags,
        "achievementUnlocked": sorted(set(achievement_ids), key=achievement_ids.index),
        "nextEventId": None,
        "nextEvent": None,
    }


def read_doc() -> tuple[list[dict], list[dict], dict]:
    document = docx.Document(DOCX_PATH)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    events: list[dict] = []
    current: dict | None = None
    intro: dict = {}
    initial_stats: dict[str, int] = {}
    initial_major_stats: dict[str, int] = {}

    for line in paragraphs:
        if line.startswith("标题："):
            intro["title"] = line.split("：", 1)[1].strip()
        elif line.startswith("开场文案："):
            intro["body"] = line.split("：", 1)[1].strip()
        elif line.startswith("系统吐槽："):
            intro["body"] = f"{intro.get('body', '')}\n\n{line.split('：', 1)[1].strip()}".strip()
        elif line.startswith("进入按钮："):
            intro["startButton"] = line.split("：", 1)[1].strip()
        elif re.match(r"^(上头值|精神电量|滤镜|绩点意志|就业幻觉|跑路冲动|市场敏感度|风险感知|证书焦虑)：", line):
            name, raw = line.split("：", 1)
            if name in GLOBAL_STAT_MAP:
                initial_stats[GLOBAL_STAT_MAP[name]] = int(raw)
            elif name in MAJOR_STAT_MAP:
                initial_major_stats[MAJOR_STAT_MAP[name]] = int(raw)

        if re.match(r"^(主线事件|随机事件)\s*\d+", line):
            if current:
                events.append(current)
            current = {}
            continue
        if current is None:
            continue
        if line.startswith("事件ID："):
            current["id"] = line.split("：", 1)[1].strip()
        elif line.startswith("学年学期："):
            raw = line.split("：", 1)[1].strip()
            current["semester"] = SEMESTER_MAP.get(raw, raw)
        elif line.startswith("事件类型："):
            raw = line.split("：", 1)[1].strip()
            current["type"] = TYPE_MAP.get(raw, raw)
        elif line.startswith("事件标题："):
            current["title"] = line.split("：", 1)[1].strip()
        elif line.startswith("事件描述："):
            current["description"] = line.split("：", 1)[1].strip()
        elif line.startswith("标签："):
            current["tags"] = split_slash(line.split("：", 1)[1])

    if current:
        events.append(current)

    event_count = len([e for e in events if e.get("type") in {"main", "major_random"}])
    expected_tables = event_count * 3
    if len(document.tables) < expected_tables:
        raise RuntimeError(f"expected at least {expected_tables} option tables, got {len(document.tables)}")

    for event_index, event in enumerate(events[:event_count]):
        options = [table_to_option(document.tables[event_index * 3 + i], event["id"], i) for i in range(3)]
        event.update(
            {
                "eventId": event["id"],
                "majorId": MAJOR_ID,
                "stage": "middle" if event.get("type") == "main" else "random",
                "weight": 1,
                "conditions": {},
                "triggerCondition": {},
                "options": options,
                "choices": options,
                "fallbackEventId": None,
                "resultText": "",
                "statChanges": {},
                "routeChanges": [],
                "tagsUnlocked": [],
                "achievementUnlocked": [],
                "nextEvent": None,
                "body": event.get("description", ""),
            }
        )

    settings = {
        "intro": intro,
        "initialStats": initial_stats,
        "initialMajorStats": initial_major_stats,
    }
    return events[:event_count], parse_endings(paragraphs), settings


def condition_from_text(text: str) -> dict:
    def leaf(name: str, op: str, value: str) -> dict:
        if name in GLOBAL_STAT_MAP:
            return {"type": "stat", "key": GLOBAL_STAT_MAP[name], "op": op, "value": int(value)}
        if name in MAJOR_STAT_MAP:
            return {"type": "majorStat", "key": MAJOR_STAT_MAP[name], "op": op, "value": int(value)}
        return {"type": "flag", "key": name}

    conditions = []
    for name, op, value in re.findall(r"([A-Za-z0-9_\u4e00-\u9fff]+)\s*(>=|<=|>|<|=)\s*(\d+)", text):
        conditions.append(leaf(name, "==" if op == "=" else op, value))
    for marker in re.findall(r"route\s*包含\s*([A-Za-z0-9_]+)", text):
        route = ROUTE_MARKER_MAP.get(marker)
        if route:
            conditions.append({"type": "route", "key": route})
        else:
            conditions.append({"type": "flag", "key": f"{MAJOR_ID}_route_marker_{marker}"})
    return conditions[0] if len(conditions) == 1 else {"all": conditions}


def parse_endings(paragraphs: list[str]) -> list[dict]:
    endings = []
    current: dict | None = None
    priority = 80
    for line in paragraphs:
        if re.match(r"^结局\s*\d+", line):
            if current:
                endings.append(current)
            title = line.split("｜", 1)[1].strip() if "｜" in line else line
            current = {"title": f"金融学·{title}", "priority": priority}
            priority -= 5
            continue
        if not current:
            continue
        if line.startswith("结局ID："):
            ending_id = line.split("：", 1)[1].strip()
            current["id"] = ending_id
            current["endingId"] = ending_id
        elif line.startswith("触发条件："):
            current["condition"] = condition_from_text(line.split("：", 1)[1].strip())
        elif line.startswith("一句话总结："):
            current["description"] = line.split("：", 1)[1].strip()
        elif line.startswith("系统诊断："):
            current["advice"] = line.split("：", 1)[1].strip()
        elif line.startswith("分享文案："):
            current["shareText"] = line.split("：", 1)[1].strip()
    if current:
        endings.append(current)

    for ending in endings:
        ending.setdefault("id", ending["title"])
        ending.setdefault("endingId", ending["id"])
        ending["majorId"] = MAJOR_ID
        ending.setdefault("description", "")
        ending.setdefault("condition", {})
        ending.setdefault("shareText", f"我在金融学副本里活成了：{ending['title']}。")
        ending.setdefault("advice", ending["description"])
    return endings


def update_major_config(new_events: list[dict], new_endings: list[dict], settings: dict) -> None:
    path = SCRIPT_DIR / "majors.json"
    majors = json.loads(path.read_text())
    for major in majors:
        if major["id"] != MAJOR_ID:
            continue
        major["intro"] = settings["intro"] or major.get("intro", {})
        major["initialStats"].update(settings["initialStats"])
        for stat in major["majorStats"]:
            if stat["key"] in settings["initialMajorStats"]:
                stat["initialValue"] = settings["initialMajorStats"][stat["key"]]
        for stat in major["majorStats"]:
            if stat["key"] == "quantPressure":
                stat["name"] = "市场敏感度"
                stat["description"] = "理解市场、价格、周期和信息变化的敏感度。"
            elif stat["key"] == "internshipCompetition":
                stat["name"] = "风险感知"
                stat["description"] = "识别风险、门槛和行业周期的能力。"
            elif stat["key"] == "certificateAnxiety":
                stat["name"] = "证书焦虑"
                stat["description"] = "证书、考试和路线选择带来的焦虑。"
        main_by_sem = {event["semester"]: event["id"] for event in new_events if event["type"] == "main"}
        for item in major["timeline"]:
            sem = item.get("key") or item.get("semester")
            if sem in main_by_sem:
                item["mainEventIds"] = [main_by_sem[sem]]
                item["theme"] = next(e["title"] for e in new_events if e["id"] == main_by_sem[sem])
        major["randomEvents"] = [event["id"] for event in new_events if event["type"] == "major_random"]
        major["achievements"] = list(ACHIEVEMENT_IDS.values())
        major["endings"] = [ending["id"] for ending in new_endings]
        break
    path.write_text(json.dumps(majors, ensure_ascii=False, indent=2) + "\n")


def update_aggregate(filename: str, major_items: list[dict]) -> None:
    path = SCRIPT_DIR / filename
    data = json.loads(path.read_text())
    data = [item for item in data if item.get("majorId") != MAJOR_ID] + major_items
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n")


def main() -> None:
    new_events, new_endings, settings = read_doc()
    existing_events = json.loads((BY_MAJOR / "finance.events.json").read_text())
    keep = [
        event
        for event in existing_events
        if event.get("type") not in {"main", "major_random"}
    ]
    finance_events = new_events + keep

    achievements = [
        {
            "id": achievement_id,
            "achievementId": achievement_id,
            "majorId": MAJOR_ID,
            "title": title,
            "description": f"在金融学副本中解锁「{title}」。",
            "condition": {"type": "flag", "key": achievement_id},
            "shareText": f"我在金融学副本里解锁了：{title}。",
        }
        for title, achievement_id in ACHIEVEMENT_IDS.items()
    ]

    (BY_MAJOR / "finance.events.json").write_text(json.dumps(finance_events, ensure_ascii=False) + "\n")
    (BY_MAJOR / "finance.endings.json").write_text(json.dumps(new_endings, ensure_ascii=False) + "\n")
    (BY_MAJOR / "finance.achievements.json").write_text(json.dumps(achievements, ensure_ascii=False) + "\n")
    update_major_config(new_events, new_endings, settings)
    update_aggregate("events.json", finance_events)
    update_aggregate("endings.json", new_endings)
    update_aggregate("achievements.json", achievements)

    print(f"imported {len(new_events)} finance events, {len(new_endings)} endings, {len(achievements)} achievements")


if __name__ == "__main__":
    main()
